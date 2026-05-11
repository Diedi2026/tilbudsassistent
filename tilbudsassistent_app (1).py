
import streamlit as st
import os, re, datetime, io
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.filterwarnings("ignore")

# ── PDF-læsning ──────────────────────────────────────────────────
def read_pdf(data: bytes) -> str:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception:
        pass
    if not text.strip():
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            for page in reader.pages:
                text += page.extract_text() or ""
        except Exception:
            pass
    return text

def read_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return ""

def read_xlsx(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                line = " ".join([str(c) for c in row if c is not None])
                if line.strip():
                    lines.append(line)
        return "\n".join(lines)
    except Exception:
        return ""

def read_uploaded(file) -> str:
    data = file.read()
    name = file.name.lower()
    if name.endswith(".pdf"):
        return read_pdf(data)
    elif name.endswith(".docx"):
        return read_docx(data)
    elif name.endswith((".xlsx", ".xls")):
        return read_xlsx(data)
    else:
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""

# ── Informationsudtræk ───────────────────────────────────────────
PRODUCT_KEYWORDS = {
    "visitkort": ["visitkort", "business card", "navnekort"],
    "flyer":     ["flyer", "løbeseddel", "reklameblad"],
    "plakat":    ["plakat", "poster", "banner"],
    "blok":      ["blok", "notesblok", "skriveblok", "notepad"],
    "folder":    ["folder", "brochure", "foldepjece"],
    "brevpapir": ["brevpapir", "brevhoved", "letterhead"],
    "kuvert":    ["kuvert", "envelope"],
    "postkort":  ["postkort", "postcard"],
    "roll-up":   ["roll-up", "rollup"],
    "tryksag":   ["tryksag", "tryk", "print"],
}
SIZE_PATTERN = re.compile(r'(\d{2,4})\s*[xX×]\s*(\d{2,4})\s*mm')
QTY_PATTERN  = re.compile(r'(\d{2,7})\s*(stk\.?|styk|eksemplarer?|oplag)', re.IGNORECASE)
DATE_PATTERN = re.compile(r'(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})')

def extract_pris(text):
    patterns = [
        re.compile(r'(?:Pris|Totalpris|pris)[:\s]+(\d[\d\s.,]*)\s*(kr\.?|DKK|dkk)', re.IGNORECASE),
        re.compile(r'\b(\d{3,6})[.,]?\d{0,2}\s*(kr\.?|DKK|dkk)', re.IGNORECASE),
    ]
    for pat in patterns:
        matches = pat.findall(text)
        if matches:
            vals = []
            for p, _ in matches:
                try:
                    clean = p.strip().replace(" ", "").replace(".", "").replace(",", ".")
                    vals.append(float(clean))
                except Exception:
                    pass
            if vals:
                return max(vals)
    return None

def extract_info(text, filename=""):
    info = {
        "filename": filename, "raw_text": text[:3000],
        "product": "ukendt", "size": None, "paper": [],
        "quantity": None, "price": None, "price_raw": None,
        "date": None, "finish": [],
    }
    tl = text.lower()
    for prod, kws in PRODUCT_KEYWORDS.items():
        if any(kw in tl for kw in kws):
            info["product"] = prod
            break
    m = SIZE_PATTERN.search(text)
    if m:
        info["size"] = f"{m.group(1)}x{m.group(2)}mm"
    for word in ["90g", "115g", "130g", "170g", "250g", "300g", "350g", "400g"]:
        if word in tl:
            info["paper"].append(word)
    for word in ["glans", "mat", "laminering", "UV", "blindpræg", "folieprint"]:
        if word.lower() in tl:
            info["finish"].append(word)
    m = QTY_PATTERN.search(text)
    if m:
        info["quantity"] = int(m.group(1).replace(".", ""))
    pris = extract_pris(text)
    if pris:
        info["price"] = pris
        info["price_raw"] = f"{pris:.2f} kr."
    m = DATE_PATTERN.search(text)
    if m:
        day, month, year = m.group(1), m.group(2), m.group(3)
        if len(year) == 2:
            year = "20" + year
        try:
            info["date"] = f"{int(day):02d}.{int(month):02d}.{year}"
        except Exception:
            pass
    return info

# ── Vidensbase ───────────────────────────────────────────────────
class TilbudsBase:
    MAX_DOCS = 200
    def __init__(self):
        self.documents = []
        self.vectorizer = None
        self.tfidf_matrix = None

    def tilfoej(self, tekst, filename="ukendt"):
        if len(self.documents) >= self.MAX_DOCS:
            return False, f"Grænse på {self.MAX_DOCS} dokumenter nået."
        if not tekst.strip():
            return False, "Tom tekst – intet at indlæse."
        info = extract_info(tekst, filename)
        self.documents.append(info)
        self._genbyg_index()
        return True, info

    def _genbyg_index(self):
        if not self.documents:
            return
        corpus = [d["raw_text"] for d in self.documents]
        self.vectorizer = TfidfVectorizer(analyzer="word", ngram_range=(1, 2), min_df=1, max_features=5000)
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)

    def soeg(self, forespørgsel, top_n=5):
        if not self.documents or self.vectorizer is None:
            return []
        q_vec = self.vectorizer.transform([forespørgsel])
        scores = cosine_similarity(q_vec, self.tfidf_matrix).flatten()
        top_idx = np.argsort(scores)[::-1][:top_n]
        return [dict(**self.documents[i], score=round(float(scores[i]), 3))
                for i in top_idx if scores[i] > 0.01]

    def antal(self):
        return len(self.documents)

    def ryd(self):
        self.__init__()

# ── Konkurrentpriser ─────────────────────────────────────────────
KONKURRENT_PRISER = {
    "visitkort": [
        {"navn": "Vistaprint",    "pris": 649,  "note": "500 stk., online standard"},
        {"navn": "Helloprint",    "pris": 589,  "note": "500 stk., bulk/online"},
        {"navn": "Flyerzone",     "pris": 695,  "note": "500 stk., hurtig levering"},
        {"navn": "Lokal trykker", "pris": 950,  "note": "Fuld service + rådgivning"},
    ],
    "flyer": [
        {"navn": "Vistaprint",    "pris": 399,  "note": "250 stk., A5"},
        {"navn": "Helloprint",    "pris": 349,  "note": "250 stk., A5"},
        {"navn": "Lokal trykker", "pris": 650,  "note": "250 stk., personlig service"},
    ],
    "plakat": [
        {"navn": "Vistaprint",    "pris": 299,  "note": "A3, enkeltsidet"},
        {"navn": "Helloprint",    "pris": 249,  "note": "A3, enkeltsidet"},
        {"navn": "Lokal trykker", "pris": 550,  "note": "A3, personlig service"},
    ],
    "blok": [
        {"navn": "Vistaprint",    "pris": 499,  "note": "50-blok, A5"},
        {"navn": "Lokal trykker", "pris": 750,  "note": "50-blok, A5 + design"},
    ],
    "folder": [
        {"navn": "Vistaprint",    "pris": 449,  "note": "250 stk., trifold"},
        {"navn": "Helloprint",    "pris": 395,  "note": "250 stk., trifold"},
        {"navn": "Lokal trykker", "pris": 850,  "note": "250 stk., fuld service"},
    ],
}

def hent_konkurrentpriser(prod):
    return KONKURRENT_PRISER.get(prod, [
        {"navn": "Vistaprint",    "pris": "–", "note": "Tjek vistaprint.dk manuelt"},
        {"navn": "Helloprint",    "pris": "–", "note": "Tjek helloprint.dk manuelt"},
        {"navn": "Lokal trykker", "pris": "–", "note": "Varierer efter opgave"},
    ])

# ── Word-eksport ─────────────────────────────────────────────────
def generer_word(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margin
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Titel
    titel = doc.add_heading("TILBUDSUDKAST", level=1)
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titel.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    doc.add_paragraph(f"⚠️  AI-genereret udkast — skal godkendes af sælger inden afsendelse")

    # Metadata tabel
    doc.add_heading("Tilbudsoplysninger", level=2)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    meta = [
        ("Tilbud nr.",  data["tilbud_nr"]),
        ("Dato",        data["dato"]),
        ("Status",      "UDKAST — afventer godkendelse"),
    ]
    for i, (k, v) in enumerate(meta):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    # Forespørgsel
    doc.add_heading("Forespørgsel", level=2)
    doc.add_paragraph(data["forespørgsel"])

    # Historiske tilbud
    doc.add_heading("Prisgrundlag — Historiske tilbud", level=2)
    htbl = doc.add_table(rows=1, cols=5)
    htbl.style = "Table Grid"
    headers = ["#", "Dokument", "Pris", "Dato", "Match %"]
    for i, h in enumerate(headers):
        htbl.rows[0].cells[i].text = h
        htbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for j, r in enumerate(data["resultater"], 1):
        row = htbl.add_row()
        row.cells[0].text = str(j)
        row.cells[1].text = r["filename"]
        row.cells[2].text = f"{r['price']:.0f} kr." if r.get("price") else "–"
        row.cells[3].text = r.get("date") or "–"
        row.cells[4].text = f"{r['score']*100:.0f}%"

    # Prisresultat
    doc.add_heading("Valgt pris", level=2)
    ptbl = doc.add_table(rows=3, cols=2)
    ptbl.style = "Table Grid"
    vp = data["valgt_pris"]
    pris_rows = [
        ("Pris ekskl. moms",  f"{vp:.2f} kr."),
        ("Moms (25%)",         f"{vp*0.25:.2f} kr."),
        ("Pris inkl. moms",    f"{vp*1.25:.2f} kr."),
    ]
    for i, (k, v) in enumerate(pris_rows):
        ptbl.rows[i].cells[0].text = k
        ptbl.rows[i].cells[1].text = v

    # Konkurrenter
    doc.add_heading("Konkurrentpriser", level=2)
    ktbl = doc.add_table(rows=1, cols=3)
    ktbl.style = "Table Grid"
    for i, h in enumerate(["Konkurrent", "Pris", "Bemærkning"]):
        ktbl.rows[0].cells[i].text = h
        ktbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for k in data["konkurrenter"]:
        row = ktbl.add_row()
        row.cells[0].text = k["navn"]
        row.cells[1].text = f"{k['pris']} kr." if isinstance(k["pris"], (int, float)) else str(k["pris"])
        row.cells[2].text = k["note"]

    # Positionering
    numeric_k = [k["pris"] for k in data["konkurrenter"] if isinstance(k["pris"], (int, float))]
    if numeric_k:
        avg_k = sum(numeric_k) / len(numeric_k)
        diff  = vp - avg_k
        retning = f"{abs(diff):.0f} kr. UNDER" if diff < 0 else (f"{diff:.0f} kr. OVER" if diff > 0 else "PÅ NIVEAU MED")
        doc.add_heading("Positionering", level=2)
        doc.add_paragraph(
            f"Jeres valgte pris på {vp:.0f} kr. er {retning} "
            f"konkurrenternes gennemsnit ({avg_k:.0f} kr.)."
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── PDF-eksport ──────────────────────────────────────────────────
def generer_pdf(data: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2.5*cm, rightMargin=2.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    BLÅ    = colors.HexColor("#1A5F7A")
    LYS    = colors.HexColor("#E8F4F8")
    MRK    = colors.HexColor("#2C3E50")

    title_style = ParagraphStyle("Title2", parent=styles["Title"],
                                 textColor=BLÅ, fontSize=20, spaceAfter=4)
    h2_style    = ParagraphStyle("H2", parent=styles["Heading2"],
                                 textColor=BLÅ, fontSize=13, spaceBefore=12, spaceAfter=4)
    body_style  = ParagraphStyle("Body2", parent=styles["Normal"],
                                 fontSize=10, leading=14)
    warn_style  = ParagraphStyle("Warn", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.HexColor("#E67E22"),
                                 leading=12)

    def tabel_style_fn(header_color=BLÅ):
        return TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0),  header_color),
            ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, 0),  9),
            ("BACKGROUND",  (0, 1), (-1, -1), LYS),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LYS]),
            ("FONTSIZE",    (0, 1), (-1, -1), 9),
            ("GRID",        (0, 0), (-1, -1), 0.5, colors.HexColor("#BDC3C7")),
            ("LEFTPADDING",  (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ])

    story = []

    # Header
    story.append(Paragraph("TILBUDSUDKAST", title_style))
    story.append(Paragraph("AI-genereret — skal godkendes af sælger inden afsendelse", warn_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLÅ, spaceAfter=8))

    # Meta
    story.append(Paragraph("Tilbudsoplysninger", h2_style))
    meta_data = [
        ["Tilbud nr.",  data["tilbud_nr"]],
        ["Dato",        data["dato"]],
        ["Status",      "UDKAST — afventer godkendelse"],
    ]
    mt = Table(meta_data, colWidths=[4*cm, 12*cm])
    mt.setStyle(TableStyle([
        ("FONTNAME",     (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1,-1), 9),
        ("BACKGROUND",   (0, 0), (-1,-1), LYS),
        ("GRID",         (0, 0), (-1,-1), 0.3, colors.HexColor("#BDC3C7")),
        ("LEFTPADDING",  (0, 0), (-1,-1), 6),
        ("TOPPADDING",   (0, 0), (-1,-1), 4),
        ("BOTTOMPADDING",(0, 0), (-1,-1), 4),
    ]))
    story.append(mt)
    story.append(Spacer(1, 0.3*cm))

    # Forespørgsel
    story.append(Paragraph("Forespørgsel", h2_style))
    story.append(Paragraph(data["forespørgsel"], body_style))
    story.append(Spacer(1, 0.3*cm))

    # Historiske tilbud
    story.append(Paragraph("Prisgrundlag — Historiske tilbud", h2_style))
    ht_data = [["#", "Dokument", "Pris", "Dato", "Match %"]]
    for j, r in enumerate(data["resultater"], 1):
        ht_data.append([
            str(j),
            r["filename"][:35],
            f"{r['price']:.0f} kr." if r.get("price") else "–",
            r.get("date") or "–",
            f"{r['score']*100:.0f}%",
        ])
    ht = Table(ht_data, colWidths=[0.8*cm, 7.5*cm, 2.5*cm, 2.5*cm, 2.0*cm])
    ht.setStyle(tabel_style_fn())
    story.append(ht)
    story.append(Spacer(1, 0.3*cm))

    # Valgt pris
    story.append(Paragraph("Valgt pris", h2_style))
    vp = data["valgt_pris"]
    pt_data = [
        ["Pris ekskl. moms",  f"{vp:.2f} kr."],
        ["Moms (25%)",         f"{vp*0.25:.2f} kr."],
        ["Pris inkl. moms",    f"{vp*1.25:.2f} kr."],
    ]
    pt = Table(pt_data, colWidths=[5*cm, 4*cm])
    pt.setStyle(TableStyle([
        ("FONTNAME",     (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1,-1), 10),
        ("BACKGROUND",   (0, 2), (-1, 2), BLÅ),
        ("TEXTCOLOR",    (0, 2), (-1, 2), colors.white),
        ("FONTNAME",     (0, 2), (-1, 2), "Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,0),(-1,1), [colors.white, LYS]),
        ("GRID",         (0, 0), (-1,-1), 0.3, colors.HexColor("#BDC3C7")),
        ("LEFTPADDING",  (0, 0), (-1,-1), 8),
        ("TOPPADDING",   (0, 0), (-1,-1), 5),
        ("BOTTOMPADDING",(0, 0), (-1,-1), 5),
    ]))
    story.append(pt)
    story.append(Spacer(1, 0.3*cm))

    # Konkurrenter
    story.append(Paragraph("Konkurrentpriser", h2_style))
    kt_data = [["Konkurrent", "Pris", "Bemærkning"]]
    for k in data["konkurrenter"]:
        kt_data.append([
            k["navn"],
            f"{k['pris']} kr." if isinstance(k["pris"], (int, float)) else str(k["pris"]),
            k["note"],
        ])
    kt = Table(kt_data, colWidths=[4.5*cm, 2.5*cm, 8.5*cm])
    kt.setStyle(tabel_style_fn())
    story.append(kt)
    story.append(Spacer(1, 0.3*cm))

    # Positionering
    numeric_k = [k["pris"] for k in data["konkurrenter"] if isinstance(k["pris"], (int, float))]
    if numeric_k:
        avg_k = sum(numeric_k) / len(numeric_k)
        diff  = vp - avg_k
        if diff < 0:
            pos_tekst = f"✅  Jeres pris på <b>{vp:.0f} kr.</b> er <b>{abs(diff):.0f} kr. UNDER</b> konkurrenternes gennemsnit ({avg_k:.0f} kr.) — attraktivt konkurrencedygtigt."
        elif diff > 0:
            pos_tekst = f"ℹ️  Jeres pris på <b>{vp:.0f} kr.</b> er <b>{diff:.0f} kr. OVER</b> konkurrenternes gennemsnit ({avg_k:.0f} kr.) — positioneret som premium/fuld service."
        else:
            pos_tekst = f"➡️  Jeres pris på <b>{vp:.0f} kr.</b> er præcis på konkurrenternes gennemsnit ({avg_k:.0f} kr.)."
        story.append(Paragraph("Positionering", h2_style))
        pos_style = ParagraphStyle("Pos", parent=body_style, backColor=LYS,
                                   borderPad=8, borderColor=BLÅ, borderWidth=1)
        story.append(Paragraph(pos_tekst, pos_style))

    doc.build(story)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════
#  STREAMLIT APP
# ══════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="Tilbudsassistent POC",
    page_icon="🖨️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ──────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1A5F7A, #2ECC71);
        color: white; padding: 1.5rem 2rem; border-radius: 12px;
        margin-bottom: 1.5rem;
    }
    .main-header h1 { color: white; margin: 0; font-size: 1.8rem; }
    .main-header p  { color: rgba(255,255,255,0.85); margin: 0.3rem 0 0; font-size: 0.95rem; }
    .metric-card {
        background: white; border: 1px solid #E0E0E0;
        border-radius: 10px; padding: 1rem 1.2rem;
        box-shadow: 0 2px 6px rgba(0,0,0,0.06); text-align: center;
    }
    .metric-card .num { font-size: 2rem; font-weight: 700; color: #1A5F7A; }
    .metric-card .lbl { font-size: 0.8rem; color: #888; margin-top: 0.2rem; }
    .pris-box {
        background: #E8F4F8; border-left: 4px solid #1A5F7A;
        border-radius: 8px; padding: 1rem 1.2rem; margin: 0.5rem 0;
    }
    .valgt-pris {
        background: linear-gradient(135deg, #1A5F7A, #148F77);
        color: white; border-radius: 10px; padding: 1.2rem;
        text-align: center; font-size: 1.6rem; font-weight: 700;
        margin: 1rem 0;
    }
    .konkurrent-under { color: #27AE60; font-weight: 600; }
    .konkurrent-over  { color: #E67E22; font-weight: 600; }
    .warning-box {
        background: #FFF3CD; border: 1px solid #FFC107;
        border-radius: 8px; padding: 0.8rem 1rem;
        font-size: 0.9rem; margin: 0.5rem 0;
    }
    div[data-testid="stSidebarContent"] { background: #F0F8FF; }
</style>
""", unsafe_allow_html=True)

# ── Session state ────────────────────────────────────────────────
if "base" not in st.session_state:
    st.session_state.base = TilbudsBase()
if "tilbud_data" not in st.session_state:
    st.session_state.tilbud_data = None
if "resultater" not in st.session_state:
    st.session_state.resultater = []

base: TilbudsBase = st.session_state.base

# ── HEADER ───────────────────────────────────────────────────────
st.markdown("""
<div class="main-header">
  <h1>🖨️ Tilbudsassistent POC</h1>
  <p>AI-drevet tilbudsgenerering fra historiske tilbud · Offline prototype · Produktion: Azure + Power Automate</p>
</div>
""", unsafe_allow_html=True)

# ── SIDEBAR ───────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📂 Vidensbase")

    c1, c2, c3 = st.columns(3)
    c1.markdown(f'<div class="metric-card"><div class="num">{base.antal()}</div><div class="lbl">Tilbud</div></div>', unsafe_allow_html=True)
    c2.markdown(f'<div class="metric-card"><div class="num">{TilbudsBase.MAX_DOCS}</div><div class="lbl">Maks.</div></div>', unsafe_allow_html=True)
    c3.markdown(f'<div class="metric-card"><div class="num">{TilbudsBase.MAX_DOCS - base.antal()}</div><div class="lbl">Ledig</div></div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 📎 Upload dokumenter")
    st.caption("Understøttede formater: PDF, DOCX, XLSX, TXT")
    uploaded_files = st.file_uploader(
        "Vælg en eller flere filer",
        type=["pdf", "docx", "xlsx", "xls", "txt"],
        accept_multiple_files=True,
        key="uploader"
    )
    if uploaded_files:
        nye = 0
        for uf in uploaded_files:
            allerede = any(d["filename"] == uf.name for d in base.documents)
            if not allerede:
                tekst = read_uploaded(uf)
                ok, info = base.tilfoej(tekst, uf.name)
                if ok:
                    nye += 1
        if nye > 0:
            st.success(f"✅ {nye} nyt/nye dokument(er) tilføjet!")
            st.rerun()

    st.markdown("---")
    st.markdown("### ✏️ Tilføj manuelt")
    with st.expander("Indsæt tilbudstekst", expanded=False):
        manual_fn   = st.text_input("Titel/filnavn", "manuelt-tilbud.txt")
        manual_text = st.text_area("Tilbudstekst", height=150,
            placeholder="Skriv eller indsæt tilbudstekst her...\nEks: Visitkort 85x55mm, 400g, 500 stk.\nPris: 875 kr.")
        if st.button("➕ Tilføj til vidensbase", use_container_width=True):
            if manual_text.strip():
                ok, info = base.tilfoej(manual_text, manual_fn)
                if ok:
                    st.success(f"✅ Tilføjet: {info['product']} | {info.get('price_raw','–')}")
                    st.rerun()
                else:
                    st.error(info)
            else:
                st.warning("Ingen tekst at tilføje.")

    st.markdown("---")
    if st.button("🗑️ Ryd vidensbase", use_container_width=True, type="secondary"):
        base.ryd()
        st.session_state.tilbud_data = None
        st.session_state.resultater  = []
        st.rerun()

# ── TABS ──────────────────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs(["🔍 Ny forespørgsel", "📋 Vidensbase", "ℹ️ Om systemet"])

# ══════════════════════════════════════════════════════════════════
#  TAB 1: NY FORESPØRGSEL
# ══════════════════════════════════════════════════════════════════
with tab1:
    if base.antal() == 0:
        st.markdown('<div class="warning-box">⚠️ <b>Vidensbasen er tom.</b> Upload historiske tilbud i sidepanelet til venstre før du sender en forespørgsel.</div>', unsafe_allow_html=True)
    else:
        st.markdown(f"### 📨 Beskriv den nye opgave")
        col_inp, col_up = st.columns([3, 1])
        with col_inp:
            forespørgsel_text = st.text_area(
                "Beskriv produktet du ønsker tilbud på:",
                height=120,
                placeholder="Eks: Visitkort 85x55mm, 400g glans laminering, 500 stk., CMYK 4/4 farver\nEller: Flyer A5, 130g mat papir, 1000 stk.\nEller: Notesblok A4, 100 blokke med logo tryk",
            )
        with col_up:
            st.markdown("**...eller upload forespørgsel:**")
            req_file = st.file_uploader("Upload fil", type=["pdf","docx","txt"], key="req_up", label_visibility="collapsed")
            if req_file:
                forespørgsel_text = read_uploaded(req_file)
                st.success(f"✅ Indlæst: {req_file.name}")

        top_n = st.slider("Antal sammenlignelige tilbud at vise:", 2, min(7, base.antal()), min(5, base.antal()))

        if st.button("🔍 Find sammenlignelige tilbud", type="primary", use_container_width=True):
            if not forespørgsel_text.strip():
                st.warning("Beskriv venligst forespørgslen først.")
            else:
                with st.spinner("Søger i vidensbasen..."):
                    resultater = base.soeg(forespørgsel_text, top_n=top_n)
                    st.session_state.resultater = resultater
                    st.session_state.tilbud_data = None
                if not resultater:
                    st.error("Ingen relevante tilbud fundet. Prøv med andre søgeord.")
                else:
                    st.success(f"✅ Fandt {len(resultater)} sammenlignelige tilbud!")

        # ── Resultater ───────────────────────────────────────────
        resultater = st.session_state.resultater
        if resultater:
            st.markdown("---")
            st.markdown("### 📊 Sammenlignelige historiske tilbud")

            # Tabel
            import pandas as pd
            df = pd.DataFrame([{
                "Rang":      i+1,
                "Dokument":  r["filename"],
                "Produkt":   r["product"],
                "Størrelse": r.get("size") or "–",
                "Antal":     str(r.get("quantity") or "–"),
                "Pris":      f"{r['price']:.0f} kr." if r.get("price") else "–",
                "Dato":      r.get("date") or "–",
                "Match":     f"{r['score']*100:.0f}%",
            } for i, r in enumerate(resultater)])
            st.dataframe(df, use_container_width=True, hide_index=True)

            # Prisstatistik
            med_pris = [r for r in resultater if r.get("price")]
            if med_pris:
                priser = [r["price"] for r in med_pris]
                gns    = sum(priser) / len(priser)
                lavest = min(priser)
                hoejest= max(priser)
                senest_r = sorted(med_pris, key=lambda x: x.get("date") or "0", reverse=True)[0]
                senest   = senest_r["price"]

                st.markdown("---")
                st.markdown("### 💰 Prisstatistik og valg")

                cols = st.columns(4)
                cols[0].metric("Gennemsnitspris", f"{gns:.0f} kr.")
                cols[1].metric("Laveste pris",    f"{lavest:.0f} kr.")
                cols[2].metric("Seneste pris",     f"{senest:.0f} kr.")
                cols[3].metric("Højeste pris",     f"{hoejest:.0f} kr.")

                st.markdown("**Vælg prisgrundlag:**")
                pris_valg = st.radio(
                    "Pristype",
                    options=["Gennemsnitspris", "Laveste pris", "Seneste pris", "Højeste pris", "Manuel pris"],
                    horizontal=True,
                    label_visibility="collapsed"
                )
                if pris_valg == "Gennemsnitspris": foreslået = gns
                elif pris_valg == "Laveste pris":  foreslået = lavest
                elif pris_valg == "Seneste pris":  foreslået = senest
                elif pris_valg == "Højeste pris":  foreslået = hoejest
                else:                               foreslået = gns

                if pris_valg == "Manuel pris":
                    valgt_pris = st.number_input("Angiv pris (kr. ekskl. moms):",
                                                  min_value=0.0, value=float(round(gns)),
                                                  step=50.0, format="%.0f")
                else:
                    valgt_pris = st.number_input("Juster pris om ønsket (kr. ekskl. moms):",
                                                  min_value=0.0, value=float(round(foreslået)),
                                                  step=50.0, format="%.0f")

                st.markdown(f'<div class="valgt-pris">Valgt pris: {valgt_pris:.0f} kr. ekskl. moms &nbsp;·&nbsp; {valgt_pris*1.25:.0f} kr. inkl. moms</div>', unsafe_allow_html=True)

                # ── Konkurrentpriser ────────────────────────────
                finfo    = extract_info(forespørgsel_text)
                produkt  = finfo["product"]
                if produkt == "ukendt" and resultater:
                    produkt = resultater[0]["product"]
                konkurrenter = hent_konkurrentpriser(produkt)

                st.markdown("---")
                st.markdown(f"### 🏆 Konkurrentpriser — *{produkt}*")

                k_cols = st.columns(len(konkurrenter))
                numeric_k = [k["pris"] for k in konkurrenter if isinstance(k["pris"], (int, float))]
                avg_k = sum(numeric_k)/len(numeric_k) if numeric_k else None

                for i, (col, k) in enumerate(zip(k_cols, konkurrenter)):
                    with col:
                        pris_str = f"{k['pris']} kr." if isinstance(k["pris"], (int, float)) else "–"
                        delta    = None
                        if avg_k and isinstance(k["pris"], (int, float)):
                            delta = f"gns: {avg_k:.0f} kr."
                        col.metric(k["navn"], pris_str, delta)
                        st.caption(k["note"])

                if avg_k:
                    diff = valgt_pris - avg_k
                    if diff < 0:
                        st.markdown(f'<p class="konkurrent-under">✅ Jeres pris er {abs(diff):.0f} kr. UNDER konkurrenternes gennemsnit ({avg_k:.0f} kr.) — konkurrencedygtigt positioneret.</p>', unsafe_allow_html=True)
                    elif diff > 0:
                        st.markdown(f'<p class="konkurrent-over">ℹ️ Jeres pris er {diff:.0f} kr. OVER konkurrenternes gennemsnit ({avg_k:.0f} kr.) — positioneret som premium/fuld service.</p>', unsafe_allow_html=True)
                    else:
                        st.markdown("➡️ Jeres pris er præcis på konkurrenternes gennemsnit.")

                # ── Generer tilbud ───────────────────────────────
                st.markdown("---")
                if st.button("📄 Generer tilbudsudkast", type="primary", use_container_width=True):
                    tilbud_nr = f"TIL-{datetime.date.today().strftime('%Y%m%d')}-001"
                    dato      = datetime.date.today().strftime("%d.%m.%Y")
                    st.session_state.tilbud_data = {
                        "tilbud_nr":    tilbud_nr,
                        "dato":         dato,
                        "forespørgsel": forespørgsel_text,
                        "resultater":   resultater,
                        "valgt_pris":   valgt_pris,
                        "konkurrenter": konkurrenter,
                        "produkt":      produkt,
                    }

            else:
                st.info("Ingen af de fundne tilbud indeholder prisinformation. Tilføj pris manuelt.")

    # ── Tilbudsudkast + download ──────────────────────────────────
    td = st.session_state.tilbud_data
    if td:
        st.markdown("---")
        st.markdown("### ✅ Tilbudsudkast")
        st.markdown('<div class="warning-box">⚠️ <b>Husk:</b> Dette er et AI-genereret udkast. Gennemgå og godkend inden afsendelse til kunde.</div>', unsafe_allow_html=True)

        with st.expander("📋 Vis tilbudsdetaljer", expanded=True):
            c1, c2, c3 = st.columns(3)
            c1.markdown(f"**Tilbud nr.:** {td['tilbud_nr']}")
            c2.markdown(f"**Dato:** {td['dato']}")
            c3.markdown(f"**Produkt:** {td['produkt']}")
            st.markdown(f"**Forespørgsel:** {td['forespørgsel']}")
            st.markdown(f"**Valgt pris:** {td['valgt_pris']:.0f} kr. ekskl. moms &nbsp;|&nbsp; {td['valgt_pris']*1.25:.0f} kr. inkl. moms")

        st.markdown("#### 📥 Download tilbudsudkast")
        dl1, dl2 = st.columns(2)
        with dl1:
            try:
                word_bytes = generer_word(td)
                st.download_button(
                    label="📝 Download som Word (.docx)",
                    data=word_bytes,
                    file_name=f"{td['tilbud_nr']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Word-eksport fejlede: {e}")

        with dl2:
            try:
                pdf_bytes = generer_pdf(td)
                st.download_button(
                    label="📄 Download som PDF",
                    data=pdf_bytes,
                    file_name=f"{td['tilbud_nr']}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF-eksport fejlede: {e}")

# ══════════════════════════════════════════════════════════════════
#  TAB 2: VIDENSBASE
# ══════════════════════════════════════════════════════════════════
with tab2:
    st.markdown("### 📋 Alle indlæste dokumenter i vidensbasen")
    if base.antal() == 0:
        st.info("Vidensbasen er tom. Upload dokumenter i sidepanelet.")
    else:
        import pandas as pd
        df2 = pd.DataFrame([{
            "#":         i+1,
            "Filnavn":   d["filename"],
            "Produkt":   d["product"],
            "Størrelse": d.get("size") or "–",
            "Papir":     ", ".join(d.get("paper", [])) or "–",
            "Finish":    ", ".join(d.get("finish", [])) or "–",
            "Antal":     str(d.get("quantity") or "–"),
            "Pris":      f"{d['price']:.0f} kr." if d.get("price") else "–",
            "Dato":      d.get("date") or "–",
        } for i, d in enumerate(base.documents)])
        st.dataframe(df2, use_container_width=True, hide_index=True)
        st.caption(f"Vidensbase: {base.antal()} / {TilbudsBase.MAX_DOCS} dokumenter")

# ══════════════════════════════════════════════════════════════════
#  TAB 3: OM SYSTEMET
# ══════════════════════════════════════════════════════════════════
with tab3:
    st.markdown("""
### 🖨️ Tilbudsassistent POC — Systeminfo

**Denne prototype demonstrerer:**
- Upload af historiske tilbud (PDF, DOCX, XLSX, TXT)
- Automatisk dataudtræk (produkt, størrelse, pris, dato)
- Semantisk søgning efter de 3–5 mest sammenlignelige tilbud
- Prisvalg: gennemsnit / laveste / seneste / manuel
- Konkurrentprissammenligning
- Eksport til Word (.docx) og PDF

---
### 🏗️ Arkitektur

| Komponent | POC (nu) | Produktion |
|---|---|---|
| Søgning | TF-IDF + cosine similarity | Azure AI Search + OpenAI embeddings |
| Sprogmodel | Regelbaseret udtræk | Azure OpenAI GPT-4o |
| Workflow | Streamlit (lokal) | Power Automate |
| Brugerflade | Streamlit | Power Apps / Teams Copilot |

---
### 📋 Kapacitet
- **Maks. dokumenter:** 200 pr. session
- **Understøttede formater:** PDF, DOCX, XLSX, TXT
- **Moms:** 25% (dansk standard)
- **Tilbudseksport:** Word + PDF

---
### ⚠️ Vigtige noter
- Vidensbasen gemmes **kun i den aktive session** — lukker du browseren, skal dokumenter genindlæses.
- Alle genererede tilbud er **udkast** og skal godkendes af en sælger.
- Konkurrentpriser er eksempeldata — opdatér `KONKURRENT_PRISER` i koden med reelle priser.
    """)
